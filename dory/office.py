import docx
import xlrd
import xlwt


class Docx:
    def __init__(self, filename=None):
        self.filename = filename
        if filename:
            self.docx = docx.Document(filename)
        else:
            self.docx = docx.Document()

    def get_paragraphs(self):
        """一维列表 所有段落"""
        paras = self.docx.paragraphs
        paras = [para.text for para in paras]
        return paras

    def get_tables(self):
        """三维列表 table row cell"""
        tables = self.docx.tables
        result = []
        for table in tables:
            result.append([])
            rows = table.rows
            for row in rows:
                result[-1].append([])
                for cell in row.cells:
                    result[-1][-1].append(cell.text)
        return result

    def add_heading(self, text, level=1):
        """添加level级标题"""
        self.docx.add_heading(text, level)

    def add_paragraph(self, text, style=None):
        """style: Normal, Heading 1, ..."""
        self.docx.add_paragraph(text, style)

    def add_page_break(self):
        """换页"""
        self.docx.add_page_break()

    def save(self, filename=None):
        if not filename:
            filename = self.filename
        self.docx.save(filename)


class Excel:
    def __init__(self, filename=None):
        self.filename = filename
        self.data = {}
        if filename:
            self._read(filename)

    def _read(self, filename):
        excel = xlrd.open_workbook(filename)
        for sheet_name in excel.sheet_names():
            sheet = excel.sheet_by_name(sheet_name)
            content = [sheet.row_values(i) for i in range(sheet.nrows)]
            self.data[sheet_name] = content

    def get_data(self):
        """{sheetname: row[cell]}"""
        return self.data

    def set_data(self, data):
        if type(data) != dict:
            self.data = {'Sheet1': data}
        else:
            self.data = data

    def __str__(self):
        return str(self.data)

    def __getitem__(self, item):
        if item in self.data:
            return self.data[item]
        else:
            self.data[item] = []
            return self.data[item]

    def __len__(self):
        return len(self.data)

    def __iter__(self):
        return iter(self.data)

    def save(self, filename, data=None):
        if data:
            self.set_data(data)
        excel = xlwt.Workbook()
        for sheet_name in self.data:
            sheet = excel.add_sheet(sheet_name)
            for i in range(len(self.data[sheet_name])):
                for j in range(len(self.data[sheet_name][i])):
                    sheet.write(i, j, self.data[sheet_name][i][j])
        excel.save(filename)


if __name__ == '__main__':
    a = Excel('a.xlsx')
    print(a['Sheet1'])
